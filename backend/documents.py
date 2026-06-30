"""PDF and DOCX ingestion and moderation helpers."""

from __future__ import annotations

import re
import tempfile
import time
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from backend.image_io import ImageInputError, validate_image_url
from backend.pipeline import pii_detector, text_safety
from backend.pipeline.text_moderation import moderate_text
from backend.reports import CATEGORY_LABELS, calculate_overall_score, decide, normalize_categories

MAX_DOCUMENT_BYTES = 25 * 1024 * 1024
MAX_PDF_PAGES = 80
MAX_DOCX_UNCOMPRESSED_BYTES = 80 * 1024 * 1024
URL_RE = re.compile(r"https?://[^\s)>'\"]+", re.IGNORECASE)
PHISHING_TERMS = {"verify", "password", "account suspended", "login", "wallet", "urgent"}
COPYRIGHT_TERMS = {"copyright", "all rights reserved", "confidential", "proprietary"}
SENSITIVE_TERMS = {"passport", "aadhaar", "ssn", "tax id", "bank statement", "invoice"}
FINANCIAL_TERMS = {"bank", "account number", "iban", "ifsc", "credit card", "payment"}
IDENTITY_TERMS = {"passport", "driver license", "national id", "aadhaar", "pan card"}
MALWARE_EXTENSIONS = (".exe", ".dll", ".bat", ".cmd", ".scr", ".ps1", ".vbs", ".js")


@dataclass(frozen=True)
class DocumentInput:
    """Validated document file prepared for moderation."""

    path: Path
    filename: str
    content_type: str | None
    size_bytes: int


class DocumentInputError(ValueError):
    """Raised when a submitted document cannot be safely processed."""


def write_document_upload(
    contents: bytes,
    filename: str,
    content_type: str | None,
    expected_suffix: str,
) -> DocumentInput:
    """Persist and validate a PDF or DOCX upload."""

    if not contents:
        raise DocumentInputError("Document upload is empty.")
    if len(contents) > MAX_DOCUMENT_BYTES:
        raise DocumentInputError("Document exceeds the 25 MB limit.")
    suffix = Path(filename or f"upload{expected_suffix}").suffix.lower()
    if suffix != expected_suffix:
        raise DocumentInputError(f"Only {expected_suffix} files are supported.")

    temp = tempfile.NamedTemporaryFile(delete=False, suffix=expected_suffix)
    try:
        temp.write(contents)
        temp.close()
        return DocumentInput(Path(temp.name), filename or temp.name, content_type, len(contents))
    except Exception:
        Path(temp.name).unlink(missing_ok=True)
        raise


def download_document(url: str, expected_suffix: str) -> DocumentInput:
    """Download a public HTTPS document with shared SSRF protections."""

    import requests

    try:
        current_url = validate_image_url(url)
        for _ in range(4):
            response = requests.get(current_url, timeout=15, stream=True, allow_redirects=False)
            if response.is_redirect or response.is_permanent_redirect:
                location = response.headers.get("location")
                if not location:
                    raise DocumentInputError("Document URL redirect is missing a target.")
                current_url = validate_image_url(requests.compat.urljoin(current_url, location))
                continue

            response.raise_for_status()
            filename = Path(current_url.split("?", 1)[0]).name or f"download{expected_suffix}"
            if Path(filename).suffix.lower() != expected_suffix:
                raise DocumentInputError(f"Document URL must end with {expected_suffix}.")

            temp = tempfile.NamedTemporaryFile(delete=False, suffix=expected_suffix)
            total = 0
            try:
                for chunk in response.iter_content(chunk_size=64 * 1024):
                    if not chunk:
                        continue
                    total += len(chunk)
                    if total > MAX_DOCUMENT_BYTES:
                        raise DocumentInputError("Remote document exceeds the 25 MB limit.")
                    temp.write(chunk)
                temp.close()
                return DocumentInput(
                    Path(temp.name), filename, response.headers.get("content-type"), total
                )
            except Exception:
                Path(temp.name).unlink(missing_ok=True)
                raise
        raise DocumentInputError("Document URL redirects too many times.")
    except ImageInputError as exc:
        raise DocumentInputError(str(exc).replace("Image URL", "Document URL")) from exc


def _link_scores(text: str, links: list[str]) -> dict[str, float]:
    lowered = text.lower()
    suspicious = any(term in lowered for term in PHISHING_TERMS)
    executable_link = any(
        link.lower().split("?", 1)[0].endswith(MALWARE_EXTENSIONS) for link in links
    )
    return {
        "phishing_score": 0.75 if links and suspicious else 0.0,
        "malware_links_score": 0.9 if executable_link else 0.0,
    }


def _document_scores(
    text: str, links: list[str], *, embedded_images: int, barcodes: int
) -> dict[str, float]:
    lowered = text.lower()
    scores = {
        **_link_scores(text, links),
        "copyright_notice_score": 0.65 if any(term in lowered for term in COPYRIGHT_TERMS) else 0.0,
        "sensitive_document_score": (
            0.75 if any(term in lowered for term in SENSITIVE_TERMS) else 0.0
        ),
        "financial_information_score": (
            0.75 if any(term in lowered for term in FINANCIAL_TERMS) else 0.0
        ),
        "identity_document_score": 0.8 if any(term in lowered for term in IDENTITY_TERMS) else 0.0,
        "watermark_score": 0.5 if "watermark" in lowered else 0.0,
        "qr_code_score": 0.65 if "qr" in lowered or barcodes else 0.0,
        "barcode_score": 0.65 if barcodes else 0.0,
    }
    if embedded_images:
        scores["embedded_image_score"] = min(1.0, embedded_images / 10)
    return scores


def _moderate_document_text(text: str) -> dict[str, float]:
    if not text.strip():
        return {}
    text_result = moderate_text(text)
    rules = text_safety.analyze_text_safety(text, None)
    pii = pii_detector.analyze_pii(text, None)
    return {**text_result.scores, **rules, **pii}


def _build_document_report(
    *,
    file_info: dict[str, Any],
    text: str,
    links: list[str],
    metadata: dict[str, Any],
    embedded_images: int = 0,
    page_count: int | None = None,
    table_count: int = 0,
    warning: str | None = None,
    started: float,
) -> dict[str, Any]:
    raw_scores = {
        **_moderate_document_text(text),
        **_document_scores(text, links, embedded_images=embedded_images, barcodes=0),
    }
    categories = normalize_categories(raw_scores)
    for key in (
        "phishing_score",
        "malware_links_score",
        "barcode_score",
        "copyright_notice_score",
        "sensitive_document_score",
        "financial_information_score",
        "identity_document_score",
    ):
        if key in raw_scores:
            public_key = key.replace("_score", "")
            categories[public_key] = round(float(raw_scores[key]) * 100, 1)
    labels = {
        **CATEGORY_LABELS,
        "phishing": "Phishing",
        "malware_links": "Malware Links",
        "barcode": "Barcodes",
        "copyright_notice": "Copyright Notices",
        "sensitive_document": "Sensitive Documents",
        "financial_information": "Financial Information",
        "identity_document": "Identity Documents",
    }
    overall = calculate_overall_score(categories, raw_scores)
    decision = decide(overall, categories)
    preview = text[:2000]
    processing_time = round(time.perf_counter() - started, 3)
    return {
        "overall_score": overall,
        "risk_level": decision.risk_level,
        "decision": decision.decision,
        "recommendation": decision.recommendation,
        "categories": categories,
        "category_labels": labels,
        "objects": [],
        "ocr_text": "",
        "extracted_text_preview": preview,
        "document": {
            "file_info": file_info,
            "page_count": page_count,
            "processing_time_seconds": processing_time,
            "metadata": metadata,
            "links": links[:50],
            "embedded_images": embedded_images,
            "table_count": table_count,
            "warning": warning,
        },
        "error": warning,
    }


def moderate_pdf(document: DocumentInput) -> dict[str, Any]:
    """Moderate a PDF document and return a report."""

    started = time.perf_counter()
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover
        raise DocumentInputError("PDF support requires pypdf to be installed.") from exc

    try:
        reader = PdfReader(str(document.path))
    except Exception as exc:
        raise DocumentInputError("PDF is corrupted or unreadable.") from exc
    if reader.is_encrypted:
        raise DocumentInputError("Password-protected PDFs are not supported.")
    if len(reader.pages) > MAX_PDF_PAGES:
        raise DocumentInputError(f"PDF exceeds the {MAX_PDF_PAGES} page limit.")

    text_parts: list[str] = []
    links: list[str] = []
    embedded_images = 0
    for page in reader.pages:
        text_parts.append(page.extract_text() or "")
        try:
            embedded_images += len(getattr(page, "images", []) or [])
        except Exception:
            pass
        annotations = page.get("/Annots") or []
        for annot in annotations:
            try:
                uri = annot.get_object().get("/A", {}).get("/URI")
                if uri:
                    links.append(str(uri))
            except Exception:
                continue

    text = "\n".join(text_parts).strip()
    links.extend(URL_RE.findall(text))
    metadata = {str(k).lstrip("/"): str(v) for k, v in (reader.metadata or {}).items()}
    warning = (
        "No machine-readable text found; scanned PDF OCR is not available in this build."
        if not text
        else None
    )
    return _build_document_report(
        file_info={
            "filename": document.filename,
            "file_type": "PDF",
            "file_size_bytes": document.size_bytes,
        },
        text=text,
        links=sorted(set(links)),
        metadata=metadata,
        embedded_images=embedded_images,
        page_count=len(reader.pages),
        warning=warning,
        started=started,
    )


def _validate_docx_zip(path: Path) -> tuple[int, list[str]]:
    try:
        with zipfile.ZipFile(path) as archive:
            infos = archive.infolist()
            total = sum(info.file_size for info in infos)
            compressed = max(sum(info.compress_size for info in infos), 1)
            if total > MAX_DOCX_UNCOMPRESSED_BYTES or total / compressed > 100:
                raise DocumentInputError("DOCX appears to be a zip bomb.")
            embedded = [info.filename for info in infos if info.filename.startswith("word/media/")]
            unsafe = [
                info.filename
                for info in infos
                if info.filename.lower().endswith(MALWARE_EXTENSIONS)
            ]
            if unsafe:
                raise DocumentInputError("DOCX contains embedded executable content.")
            return len(embedded), [info.filename for info in infos]
    except zipfile.BadZipFile as exc:
        raise DocumentInputError("DOCX is corrupted or unreadable.") from exc


def moderate_docx(document: DocumentInput) -> dict[str, Any]:
    """Moderate a DOCX document and return a report."""

    started = time.perf_counter()
    embedded_images, _ = _validate_docx_zip(document.path)
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        raise DocumentInputError("DOCX support requires python-docx to be installed.") from exc

    try:
        doc = Document(str(document.path))
    except Exception as exc:
        raise DocumentInputError("DOCX is corrupted or unreadable.") from exc

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                tables.append(" | ".join(cells))
    text = "\n".join([*paragraphs, *tables]).strip()
    links = URL_RE.findall(text)
    core = doc.core_properties
    metadata = {
        "author": core.author or "",
        "title": core.title or "",
        "subject": core.subject or "",
        "keywords": core.keywords or "",
        "created": core.created.isoformat() if core.created else "",
        "modified": core.modified.isoformat() if core.modified else "",
    }
    return _build_document_report(
        file_info={
            "filename": document.filename,
            "file_type": "DOCX",
            "file_size_bytes": document.size_bytes,
        },
        text=text,
        links=sorted(set(links)),
        metadata=metadata,
        embedded_images=embedded_images,
        table_count=len(doc.tables),
        warning=None if text else "No text found in DOCX.",
        started=started,
    )
